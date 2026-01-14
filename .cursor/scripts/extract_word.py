"""
Extract text and embedded images from Word documents.

Outputs:
  <name>_analysis/
    text_content.md   # extracted text
    images/
      image_01.png    # extracted embedded images
      ...
"""

import argparse
import os
from pathlib import Path
from typing import List, Tuple

from docx import Document


def ensure_dir(path: Path) -> Path:
    path.mkdir(parents=True, exist_ok=True)
    return path


def extract_text(doc: Document) -> str:
    """Extract paragraphs and simple tables into markdown-friendly text."""
    lines: List[str] = []

    # paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)
    # simple tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text:
                lines.append(row_text)

    return "\n\n".join(lines)


def extract_images(doc: Document, images_dir: Path) -> List[Path]:
    """Extract embedded images from docx related parts."""
    ensure_dir(images_dir)
    image_paths: List[Path] = []
    counter = 1

    for rel in doc.part.related_parts.values():
        ctype = rel.content_type
        if not ctype.startswith("image/"):
            continue
        ext = ctype.split("/")[-1]
        name = f"image_{counter:02d}.{ext}"
        out_path = images_dir / name
        with open(out_path, "wb") as f:
            f.write(rel.blob)
        image_paths.append(out_path)
        counter += 1
        print(f"[word-img] extracted -> {out_path}")

    return image_paths


def process_word(source: Path, target_dir: Path) -> Tuple[Path, Path, List[Path]]:
    analysis_dir = target_dir / f"{source.stem}_analysis"
    ensure_dir(analysis_dir)
    images_dir = ensure_dir(analysis_dir / "images")

    doc = Document(str(source))

    text = extract_text(doc)
    text_md_path = analysis_dir / "text_content.md"
    text_md_path.write_text(text, encoding="utf-8")
    print(f"[word-text] {text_md_path}")

    image_paths = extract_images(doc, images_dir)

    return analysis_dir, text_md_path, image_paths


def main():
    parser = argparse.ArgumentParser(description="Extract text + images from Word (docx/doc).")
    parser.add_argument("source", help="Word file path")
    parser.add_argument("-t", "--target-dir", default=".", help="Output base directory")
    args = parser.parse_args()

    src = Path(args.source).expanduser().resolve()
    if not src.exists():
        raise SystemExit(f"Source not found: {src}")

    target_dir = Path(args.target_dir).expanduser().resolve()
    ensure_dir(target_dir)

    analysis_dir, text_md_path, image_paths = process_word(src, target_dir)
    print(f"Analysis directory: {analysis_dir}")
    print(f"Text markdown: {text_md_path}")
    print(f"Extracted images: {len(image_paths)}")


if __name__ == "__main__":
    main()

